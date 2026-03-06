"""Content extractors for different file types.

Each extractor takes a file path and returns extracted text content.
Supports: PDF, images, Word documents, HTML.
"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Minimum chars per page to consider text extraction successful (vs image-heavy)
MIN_CHARS_PER_PAGE = 100

# Maximum pages for vision API fallback on PDFs
MAX_VISION_PAGES = 20


def extract_content(file_path: str, file_type: str, mime_type: str) -> Optional[dict]:
    """Dispatch to the appropriate extractor based on file type.

    Args:
        file_path: Path to the stored file.
        file_type: One of 'pdf', 'image', 'word', 'html'.
        mime_type: Original MIME type string.

    Returns:
        {"text": str, "page_range": Optional[str]} or None on failure.
    """
    extractors = {
        "pdf": extract_pdf,
        "image": extract_image,
        "word": extract_word,
        "html": extract_html,
    }
    extractor = extractors.get(file_type)
    if not extractor:
        logger.warning("No extractor for file type: %s", file_type)
        return None

    try:
        return extractor(file_path, mime_type)
    except Exception:
        logger.exception("Extraction failed for %s (type=%s)", file_path, file_type)
        return None


def extract_pdf(file_path: str, mime_type: str = "") -> Optional[dict]:
    """Extract text from PDF using PyMuPDF.

    Strategy:
    - Try text extraction first
    - If text is sparse (< 100 chars/page), fall back to vision API
    - For large PDFs, vision on first 20 pages only
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF (fitz) not installed. Cannot extract PDF.")
        return None

    try:
        doc = fitz.open(file_path)
    except Exception:
        logger.exception("Failed to open PDF: %s", file_path)
        return None

    page_count = len(doc)
    all_text = []
    sparse_pages = []

    for i, page in enumerate(doc):
        text = page.get_text("text").strip()
        all_text.append(text)
        if len(text) < MIN_CHARS_PER_PAGE:
            sparse_pages.append(i)

    doc.close()

    full_text = "\n\n".join(all_text)

    # If most pages are sparse, try vision API for those pages
    if len(sparse_pages) > page_count * 0.5 and _vision_available():
        vision_pages = sparse_pages[:MAX_VISION_PAGES]
        vision_text = _extract_pdf_pages_via_vision(file_path, vision_pages)
        if vision_text:
            # Replace sparse page text with vision results
            for page_num, text in zip(vision_pages, vision_text):
                if text and len(text) > len(all_text[page_num]):
                    all_text[page_num] = text
            full_text = "\n\n".join(all_text)

    page_range = "1-{}".format(page_count) if page_count > 1 else "1"

    return {"text": full_text, "page_range": page_range}


def extract_image(file_path: str, mime_type: str = "") -> Optional[dict]:
    """Extract content from an image using Claude vision API."""
    if not _vision_available():
        return {"text": "[Image uploaded but vision API not configured]"}

    try:
        import anthropic
        import base64

        image_data = Path(file_path).read_bytes()
        b64_data = base64.b64encode(image_data).decode("utf-8")

        # Determine media type
        media_type = mime_type or "image/jpeg"

        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": b64_data,
                            },
                        },
                        {
                            "type": "text",
                            "text": (
                                "Extract all text and describe all visual content "
                                "in this image. Include any data from charts, tables, "
                                "or diagrams. Format as structured text."
                            ),
                        },
                    ],
                }
            ],
        )

        text = response.content[0].text if response.content else ""
        return {"text": text}

    except Exception:
        logger.exception("Image extraction via vision API failed")
        return {"text": "[Image processing failed]"}


def extract_word(file_path: str, mime_type: str = "") -> Optional[dict]:
    """Extract text from Word (.docx) documents using python-docx.

    Preserves structure as markdown: headings, paragraphs, tables.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Cannot extract Word documents.")
        return None

    try:
        doc = Document(file_path)
    except Exception:
        logger.exception("Failed to open Word document: %s", file_path)
        return None

    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                parts.append("# {}".format(text))
            elif "Heading 2" in style_name:
                parts.append("## {}".format(text))
            elif "Heading 3" in style_name:
                parts.append("### {}".format(text))
            elif "List" in style_name:
                parts.append("- {}".format(text))
            else:
                parts.append(text)

        elif tag == "tbl":
            # Table
            for table in doc.tables:
                if table._element is element:
                    parts.append(_table_to_markdown(table))
                    break

    return {"text": "\n\n".join(parts)}


def extract_html(file_path: str, mime_type: str = "") -> Optional[dict]:
    """Extract content from HTML using trafilatura for boilerplate removal."""
    try:
        import trafilatura
    except ImportError:
        logger.warning("trafilatura not installed, using basic HTML extraction")
        return _extract_html_basic(file_path)

    try:
        html_content = Path(file_path).read_text(encoding="utf-8", errors="replace")
        text = trafilatura.extract(
            html_content,
            include_tables=True,
            include_links=True,
            output_format="txt",
        )

        if not text:
            return _extract_html_basic(file_path)

        return {"text": text}

    except Exception:
        logger.exception("HTML extraction failed for %s", file_path)
        return _extract_html_basic(file_path)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _vision_available() -> bool:
    """Check if the Anthropic API key is configured for vision calls."""
    return bool(os.environ.get("ANTHROPIC_API_KEY"))


def _extract_pdf_pages_via_vision(file_path: str, page_numbers: list[int]) -> list[str]:
    """Use Claude vision API to extract content from specific PDF pages.

    Renders pages to images and sends to the vision API.
    """
    try:
        import fitz
        import anthropic
        import base64

        doc = fitz.open(file_path)
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        results = []

        for page_num in page_numbers:
            if page_num >= len(doc):
                results.append("")
                continue

            page = doc[page_num]
            # Render page to PNG at 150 DPI
            pix = page.get_pixmap(dpi=150)
            img_data = pix.tobytes("png")
            b64_data = base64.b64encode(img_data).decode("utf-8")

            try:
                response = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=1500,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "image",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "image/png",
                                        "data": b64_data,
                                    },
                                },
                                {
                                    "type": "text",
                                    "text": "Extract all text from this document page.",
                                },
                            ],
                        }
                    ],
                )
                text = response.content[0].text if response.content else ""
                results.append(text)
            except Exception:
                logger.warning("Vision extraction failed for page %d", page_num)
                results.append("")

        doc.close()
        return results

    except Exception:
        logger.exception("PDF vision extraction failed")
        return []


def _table_to_markdown(table) -> str:
    """Convert a python-docx table to markdown format."""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| {} |".format(" | ".join(cells)))
        if i == 0:
            # Header separator
            rows.append("| {} |".format(" | ".join("---" for _ in cells)))
    return "\n".join(rows)


def _extract_html_basic(file_path: str) -> Optional[dict]:
    """Basic HTML text extraction without trafilatura."""
    try:
        import re

        html = Path(file_path).read_text(encoding="utf-8", errors="replace")
        # Strip tags
        text = re.sub(r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return {"text": text} if text else None
    except Exception:
        logger.exception("Basic HTML extraction failed")
        return None
