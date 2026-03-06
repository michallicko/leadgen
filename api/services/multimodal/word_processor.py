"""Word document (.docx) processing (BL-266).

Uses python-docx to extract paragraphs, tables, and headings,
converting them to markdown format.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class WordExtractionResult:
    """Result of Word document extraction."""

    markdown: str = ""
    headings: list[str] = field(default_factory=list)
    paragraph_count: int = 0
    table_count: int = 0
    word_count: int = 0
    error: Optional[str] = None


def extract_from_file(file_path: str) -> WordExtractionResult:
    """Extract content from a .docx file.

    Args:
        file_path: Path to the Word document.

    Returns:
        WordExtractionResult with markdown content.
    """
    try:
        with open(file_path, "rb") as f:
            return extract_from_bytes(f.read())
    except FileNotFoundError:
        return WordExtractionResult(error="File not found: {}".format(file_path))
    except Exception as exc:
        return WordExtractionResult(error="Failed to read file: {}".format(str(exc)))


def extract_from_bytes(docx_bytes: bytes) -> WordExtractionResult:
    """Extract content from .docx bytes.

    Args:
        docx_bytes: Raw .docx file content.

    Returns:
        WordExtractionResult with markdown content.
    """
    try:
        import docx
    except ImportError:
        return WordExtractionResult(
            error="python-docx not installed — run: pip install python-docx"
        )

    result = WordExtractionResult()

    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        return WordExtractionResult(
            error="Failed to parse document: {}".format(str(exc))
        )

    parts = []

    # Process paragraphs
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        result.paragraph_count += 1
        style_name = (para.style.name or "").lower() if para.style else ""

        if "heading 1" in style_name:
            parts.append("# {}".format(text))
            result.headings.append(text)
        elif "heading 2" in style_name:
            parts.append("## {}".format(text))
            result.headings.append(text)
        elif "heading 3" in style_name:
            parts.append("### {}".format(text))
            result.headings.append(text)
        elif "heading" in style_name:
            parts.append("#### {}".format(text))
            result.headings.append(text)
        elif "list" in style_name or "bullet" in style_name:
            parts.append("- {}".format(text))
        else:
            parts.append(text)

    # Process tables
    for table in document.tables:
        result.table_count += 1
        md_table = _table_to_markdown(table)
        if md_table:
            parts.append(md_table)

    result.markdown = "\n\n".join(parts)
    result.word_count = len(result.markdown.split())

    return result


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to markdown."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Build markdown table
    header = rows[0]
    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")

    for row in rows[1:]:
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[: len(header)]) + " |")

    return "\n".join(lines)
