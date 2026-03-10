"""Tests for Word document processor (BL-266)."""

import io

import pytest


class TestWordExtraction:
    def test_extract_from_bytes_no_docx(self):
        """When python-docx is not installed, returns error."""
        from api.services.multimodal.word_processor import extract_from_bytes

        # Feed invalid bytes — should get a parse error, not crash
        result = extract_from_bytes(b"not a docx file")
        assert result.error is not None

    def test_extract_from_file_not_found(self):
        from api.services.multimodal.word_processor import extract_from_file

        result = extract_from_file("/nonexistent/path.docx")
        assert result.error is not None
        assert "not found" in result.error.lower() or "Failed" in result.error

    def test_extract_real_docx(self):
        """Test with a real .docx created in-memory."""
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")

        from api.services.multimodal.word_processor import extract_from_bytes

        # Create a minimal .docx in memory
        doc = docx.Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("Another paragraph here.")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Foo"
        table.rows[1].cells[1].text = "100"

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = extract_from_bytes(docx_bytes)
        assert result.error is None
        assert "# Test Heading" in result.markdown
        assert "## Section Two" in result.markdown
        assert "test paragraph" in result.markdown
        assert result.paragraph_count >= 2
        assert result.table_count == 1
        assert "Foo" in result.markdown
        assert "Test Heading" in result.headings
        assert "Section Two" in result.headings
